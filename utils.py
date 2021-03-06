import os
import base64
import time

from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from flask_mail import Message
from flask import render_template
from extensions import mail
from flask import current_app
from flask import request

UPLOAD_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'uploads')

def save_signature(base64_str, name_3_emp_name, frm_name):
    path = os.path.join(UPLOAD_DIR, name_3_emp_name)
    file_name = '{}_{}.png'.format(path, frm_name, time.time())
    image = base64.b64decode(base64_str.split(',')[1])
    with open(file_name, 'wb') as f:
        f.write(image)
        f.close()
    return file_name

def send_link_as_mail(**kwargs):
    subject = 'PRD Form - {}'.format(kwargs['emp_name'])

    msg = Message(subject, sender='pkaur@aigbusiness.com', recipients=[
        kwargs['rev_email']
    ])

    msg.html = """Please click on the link below to sign PRD form.<br>
    <a href="http://{0}/document/{1}/{2}">Click here</a>
    """.format(request.host,kwargs['emp_code'],kwargs['reviewer_code'])

    mail.send(msg)


def send_manager_link_as_mail(**kwargs):
    subject = 'Final Rating'

    msg = Message(subject, sender='pkaur@aigbusiness.com', recipients=[
        kwargs['rev_email1']
    ])

    msg.html = """Please click on the link below to sign your PRD Form.<br>
    <a href="http://{0}/final_form/{1}/{2}">Click here</a>
    """.format(request.host,kwargs['emp_code1'],kwargs['reviewer_code1'])

    mail.send(msg)


def save_document_as_docx(**kwargs):
    document = Document()
    style = document.styles['Normal']
    font = style.font
    font.size = Pt(10)
    document.add_heading('PRD FORM')
    document.add_paragraph('Employee Information')

   

    table = document.add_table(rows=1, cols=5)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Employee Information'

    row_cells = table.add_row().cells
    row_cells[0].text = 'Employee Name :'
    row_cells[1].text = kwargs['emp_name']
    row_cells[2].text = 'Outstanding'
    row_cells[3].text = '91-100'
    row_cells[4].text = 'S'

    row_cells = table.add_row().cells
    row_cells[0].text = 'Employee Email :'
    row_cells[1].text = kwargs['emp_email']
    row_cells[2].text = 'Exceeds Expectations'
    row_cells[3].text = '81-90'
    row_cells[4].text = 'M'

    row_cells = table.add_row().cells
    row_cells[0].text = 'Employee Code :'
    row_cells[1].text = kwargs['emp_code']
    row_cells[2].text = 'Meets Expectations'
    row_cells[3].text = '71-80'
    row_cells[4].text = 'A'

    row_cells = table.add_row().cells
    row_cells[0].text = 'Job Function :'
    row_cells[1].text = kwargs['job_function']
    row_cells[2].text = 'Improvement Needed'
    row_cells[3].text = '51-70'
    row_cells[4].text = 'R'

    row_cells = table.add_row().cells
    row_cells[0].text = 'Date of Review :'
    row_cells[1].text = kwargs['date'].strftime('%d-%m-%Y')
    row_cells[2].text = 'Unacceptable'
    row_cells[3].text = 'Below 50'
    row_cells[4].text = 'T'

    row_cells = table.add_row().cells
    row_cells[0].text = 'Reviewed By :'
    row_cells[1].text = kwargs['reviewer_name']

    row_cells = table.add_row().cells
    row_cells[0].text = 'Reviewer\'s Employee Code :'
    row_cells[1].text = kwargs['reviewer_code']


    table = document.add_table(rows=1, cols=6)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Quality of Work (Effective)'

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Work is performed accurately & neatly./' \
                        'Work is consistent, thorough and complete.'
    hdr_cells[1].text = 'Self Assessment'
    hdr_cells[2].text = 'Manager Assessment'
    hdr_cells[3].text = 'Comments'
    hdr_cells[4].text = 'Total Score'
    hdr_cells[5].text = 'Achieved Score'

    row_cells = table.add_row().cells
    row_cells[0].text = kwargs['self_assessment1_comment1']
    row_cells[1].text = kwargs['self_assessment1']
    row_cells[2].text = kwargs['manager_assessment1']
    row_cells[3].text = kwargs['manager_assessment1_comment1']
    row_cells[4].text = '20'
    row_cells[5].text = kwargs['achieved_score1']

    table = document.add_table(rows=1, cols=6)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Quantity of Work (Productivity)'

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Amount of work performed on daily basis is appropriate for job function'
    hdr_cells[1].text = 'Self Assessment'
    hdr_cells[2].text = 'Manager Assessment'
    hdr_cells[3].text = 'Comments'
    hdr_cells[4].text = 'Total Score'
    hdr_cells[5].text = 'Achieved Score'

    row_cells = table.add_row().cells
    row_cells[0].text = kwargs['self_assessment2_comment2']
    row_cells[1].text = kwargs['self_assessment2']
    row_cells[2].text = kwargs['manager_assessment2']
    row_cells[3].text = kwargs['manager_assessment2_comment2']
    row_cells[4].text = '20'
    row_cells[5].text = kwargs['achieved_score2']

    table = document.add_table(rows=1, cols=6)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Job Function (Efficiency/Timeliness)'

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Understands the job requirement and has specific content knowledge /' \
                        'where appropriate, completes work within timelines & is efficient for the work'
    hdr_cells[1].text = 'Self Assessment'
    hdr_cells[2].text = 'Manager Assessment'
    hdr_cells[3].text = 'Comments'
    hdr_cells[4].text = 'Total Score'
    hdr_cells[5].text = 'Achieved Score'

    row_cells = table.add_row().cells
    row_cells[0].text = kwargs['self_assessment3_comment3']
    row_cells[1].text = kwargs['self_assessment3']
    row_cells[2].text = kwargs['manager_assessment3']
    row_cells[3].text = kwargs['manager_assessment3_comment3']
    row_cells[4].text = '30'
    row_cells[5].text = kwargs['achieved_score3']

    table = document.add_table(rows=1, cols=6)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Attendance & Punctuality'

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Is punctual and do not take uninformed leaves'
    hdr_cells[1].text = 'Self Assessment'
    hdr_cells[2].text = 'Manager Assessment'
    hdr_cells[3].text = 'Comments'
    hdr_cells[4].text = 'Total Score'
    hdr_cells[5].text = 'Achieved Score'

    row_cells = table.add_row().cells
    row_cells[0].text = kwargs['self_assessment4_comment4']
    row_cells[1].text = kwargs['self_assessment4']
    row_cells[2].text = kwargs['manager_assessment4']
    row_cells[3].text = kwargs['manager_assessment4_comment4']
    row_cells[4].text = '10'
    row_cells[5].text = kwargs['achieved_score4']

    table = document.add_table(rows=1, cols=6)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Behaviour'

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Behaviour with colleagues, Subordinates, Supervisor'
    hdr_cells[1].text = 'Self Assessment'
    hdr_cells[2].text = 'Manager Assessment'
    hdr_cells[3].text = 'Comments'
    hdr_cells[4].text = 'Total Score'
    hdr_cells[5].text = 'Achieved Score'

    row_cells = table.add_row().cells
    row_cells[0].text = kwargs['self_assessment5_comment5']
    row_cells[1].text = kwargs['self_assessment5']
    row_cells[2].text = kwargs['manager_assessment5']
    row_cells[3].text = kwargs['manager_assessment5_comment5']
    row_cells[4].text = '10'
    row_cells[5].text = kwargs['achieved_score5']

    table = document.add_table(rows=1, cols=6)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Initiatives'

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Is employee always keen on taking intitiatives  to follow the mission and values of the company'
    hdr_cells[1].text = 'Self Assessment'
    hdr_cells[2].text = 'Manager Assessment'
    hdr_cells[3].text = 'Comments'
    hdr_cells[4].text = 'Total Score'
    hdr_cells[5].text = 'Achieved Score'

    row_cells = table.add_row().cells
    row_cells[0].text = kwargs['self_assessment6_comment6']
    row_cells[1].text = kwargs['self_assessment6']
    row_cells[2].text = kwargs['manager_assessment6']
    row_cells[3].text = kwargs['manager_assessment6_comment6']
    row_cells[4].text = '10'
    row_cells[5].text = kwargs['achieved_score6']

    row_cells = table.add_row().cells
    row_cells[0].text = ''
    row_cells[1].text = ''
    row_cells[2].text = ''
    row_cells[3].text = ''
    row_cells[4].text = '100'
    row_cells[5].text = kwargs['total_achieved_score']

    table = document.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Final Manager Comments'
    hdr_cells[1].text = 'Final Manager Rating'

    row_cells = table.add_row().cells
    row_cells[0].text = kwargs['manager_final_comment']
    row_cells[1].text = kwargs['manager_final_rating']


    document.add_paragraph("""Reviewer's Signature : """)
    document.add_picture(kwargs['signature'], height=Inches(2.0))

    document.add_paragraph("""Employee Signature : """)
    document.add_picture(kwargs['signature1'], height=Inches(2.0))

    file_name_final = '%s_%s.docx' % (kwargs['emp_name'],kwargs['emp_code'])
    document.save(
        os.path.join(UPLOAD_DIR, file_name_final)
    )
    return file_name_final


def send_document_as_mail(**kwargs):
    subject = 'PRD Form - {}'.format(kwargs['emp_name'])

    msg = Message(subject, sender='pkaur@aigbusiness.com', recipients=[
        'pkaur@aigbusiness.com'
    ])

    msg.html = """Please find the attached form."""
    with current_app.open_resource(
        os.path.join(UPLOAD_DIR, kwargs['file_name'])
    ) as fp:
        msg.attach(
            filename=kwargs['file_name'],
            data=fp.read(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    mail.send(msg)